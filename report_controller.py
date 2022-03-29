import docx
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from PIL import Image


class ReportController:
    def __init__(self):
        self.doc = docx.Document()
        self._image_data = []
        self._report_title = ""

        # 調整文件左右上下邊界至 1.27 cm
        section = self.doc.sections[0]
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)

        # 設定字型及大小
        style = self.doc.styles["Normal"]
        self.doc.styles["Normal"].font.name = "標楷體"
        self.doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
        self.doc.styles["Normal"].font.size = Pt(20)

    def set_data(self, image_data, title):
        self._image_data = image_data
        self._report_title = title

    def add_table(self, start_index, count):
        # 標題
        title = self.doc.add_paragraph(self._report_title)
        title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.style = self.doc.styles["Normal"]

        # 建立表格
        table = self.doc.add_table(6, 4)
        table.style = "Table Grid"
        table.autofit = False

        # 表格全部設為垂直置中
        for i in range(6):
            for j in range(4):
                table.cell(i, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 設定表格寬度
        widths = [2.6, 9.0, 2.6]
        for i in range(3):
            for cell in table.columns[i].cells:
                cell.width = Cm(widths[i])

        # 設定表格高度
        heights = [8.8, 0.5, 1.3, 8.8, 0.5, 1.3]
        for i, row in enumerate(table.rows):
            row.height = Cm(heights[i])

        # 合併格子
        for i in [0, 3]:
            row = table.rows[i]
            left_cell = row.cells[0]
            for j in range(1, 4):
                right_cell = row.cells[j]
                left_cell.merge(right_cell)
        for i in [2, 5]:
            row = table.rows[i]
            left_cell = row.cells[1]
            for j in range(2, 4):
                right_cell = row.cells[j]
                left_cell.merge(right_cell)

        # 填上預設文字
        for i in [1, 4]:
            row = table.rows[i]

            word_position = row.cells[0].paragraphs[0]
            run = word_position.add_run("時間")
            run.font.size = Pt(12)
            word_position.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE

            word_position = row.cells[2].paragraphs[0]
            run = word_position.add_run("照片編號")
            run.font.size = Pt(12)
            word_position.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE

        for i in [2, 5]:
            row = table.rows[i]

            word_position = row.cells[0].paragraphs[0]
            run = word_position.add_run("說明")
            run.font.size = Pt(12)
            word_position.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE

            word_position = row.cells[1].paragraphs[0]
            run = word_position.add_run(" ")
            run.font.size = Pt(12)

        # 填上內容
        for i in range(count):
            # 貼圖片
            pic_position = table.rows[i * 3].cells[0].paragraphs[0]
            pic_position.add_run().add_picture(
                self._image_data[start_index + i]["file_path"], height=Cm(8.5)
            )
            pic_position.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 加時間
            datatime = self.get_datetime(self._image_data[start_index + i])
            word_position = table.rows[i * 3 + 1].cells[1].paragraphs[0]
            run = word_position.add_run(datatime)
            run.font.size = Pt(12)

            # 加編號
            word_position = table.rows[i * 3 + 1].cells[3].paragraphs[0]
            run = word_position.add_run(str(start_index + 1 + i))
            run.font.size = Pt(12)

    def generate_doc(self):
        num_iamge_data = len(self._image_data)
        num_pages = (num_iamge_data + 1) // 2

        for i in range(num_pages):
            start_index = i * 2
            if start_index + 2 <= num_iamge_data:
                self.add_table(start_index, 2)
            else:
                self.add_table(start_index, 1)
            if i != num_pages - 1:
                self.doc.add_page_break()

    def save(self, file):
        self.doc.save(file)

    def clear(self):
        self.doc = docx.Document()

    def get_datetime(self, image_data):
        # get datetime in string ex:108年09月29日17時12分17秒
        if image_data["use_image_time"]:
            img = Image.open(image_data["file_path"])
            try:
                exif = img._getexif()
                datatime = exif.get(36867)
                date = datatime.split(" ")[0].split(":")
                time = datatime.split(" ")[1].split(":")
                str_datetime = "{}年{}月{}日{}時{}分{}秒".format(
                    str(int(date[0]) - 1911),
                    date[1],
                    date[2],
                    time[0],
                    time[1],
                    time[2],
                )
            except:
                str_datetime = "年月日時分秒"
        else:
            year = image_data["time"]["year"]
            month = image_data["time"]["month"]
            day = image_data["time"]["day"]
            hour = image_data["time"]["hour"]
            minute = image_data["time"]["minute"]
            second = image_data["time"]["second"]
            str_datetime = f"{year}年{month}月{day}日{hour}時{minute}分{second}秒"

        return str_datetime
